from docx import Document
import os
import pandas as pd

def extract_and_classify_content_v5(doc_path):
    doc = Document(doc_path)
    
    # 将整个文档内容转换为单一字符串
    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"

    categories_starts = {
        "基本信息": ["摘要","患者", "姓名", "年龄", "住院号", "影像号"],
        "主诉": ["主诉"],
        "现病史": ["现病史"],
        "既往史": ["既往史", "既往"],
        "个人史": ["个人史"],
        "家族史": ["家族史"],
        "查体": ["查体", "入院查体", "体格检查"],
        "辅助检查": ["辅助检查",'术前头颅MRI','头颅MRI','头MRI','MRI','MR','PET','VEEG','EEG','胸部正位片','视频脑电图','脑电图','脑磁图','头颅CT','CT']
    }
    
    # 初始化分类内容存储字典
    categorized_content = {category: "" for category in categories_starts}
    
    # 当前分类
    current_category = None
    start_index = 0

    for category, keywords in categories_starts.items():
        for keyword in keywords:
            index = full_text.find(keyword, start_index)
            if index != -1:
                if current_category is not None:
                    # 将之前的文本归类到最近匹配的类别中
                    categorized_content[current_category] += full_text[start_index:index] + "$"
                start_index = index
                current_category = category
                break  # 匹配到一个关键词即跳出循环，继续下一个类别的匹配
                
    # 将最后一个类别的文本添加到分类中
    if current_category is not None:
        categorized_content[current_category] += full_text[start_index:] + "$"
    
    return categorized_content

def process_folder(folder_path, output_excel_path):
    results = pd.DataFrame()

    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            doc_path = os.path.join(folder_path, filename)
            content = extract_and_classify_content_v5(doc_path)
            
            current_results = pd.DataFrame([content])
            current_results['文件名'] = filename
            
            results = pd.concat([results, current_results], ignore_index=True)
    
    with pd.ExcelWriter(output_excel_path, engine='openpyxl') as writer:
        results.to_excel(writer, index=False)


if __name__ == '__main__':
    folder_path = r'E:\病例整理\先兆\病历摘要\病历摘要（未融合）'
    output_excel_path = r'D:\python code\sort_docx\files\分类结果v2.xlsx'
    process_folder(folder_path, output_excel_path)
