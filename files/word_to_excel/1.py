import pandas as pd
from docx import Document

def word_table_to_excel(doc_path, excel_path):
    # 加载Word文档
    doc = Document(doc_path)
    
    # 创建一个pandas的Excel写入器
    writer = pd.ExcelWriter(excel_path, engine='openpyxl')
    #print(type(doc.tables))
    # 遍历文档中的所有表格
    for i, table in enumerate(doc.tables):
        # 读取表格数据到一个二维列表
        data = []
        print(i+1)
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
            for cell in row.cells:
                
                print(cell.text)    
        # 将数据转换成pandas DataFrame
        df = pd.DataFrame(data)
        
        # 写入Excel文件的不同工作表
        
        sheet_name = f'Table_{i+1}'  # 工作表名基于表格的顺序
        #df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    # 保存Excel文件
    #writer._save()

# 定义Word文档和目标Excel文件的路径
doc_path = r'D:\python code\塔吉科研\word_to_excel\452108-appendix-table-2-influencing-factors-of-dolii.docx'  # 替换为你的Word文档路径
excel_path = r'D:\python code\塔吉科研\word_to_excel\1.xlsx'  # 替换为你想保存的Excel文件路径

# 执行函数
word_table_to_excel(doc_path, excel_path)
